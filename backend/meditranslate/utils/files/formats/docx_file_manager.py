import tempfile
from meditranslate.utils.files.formats.file_manager import FileManager
from io import BytesIO
import docx
from uuid import uuid4


class DocxFileManager(FileManager):
    def extract_text(self, file_stream: BytesIO) -> str:
        temp_file_name = str(uuid4())
        # with open('{temp_file_name}.docx', 'rb') as f:
        #     source_stream = StringIO(f.read())
        # document = Document(source_stream)
        # # Create a temporary file to store the .docx content
        # with tempfile.NamedTemporaryFile(delete=True, mode='wb') as temp_file:
        #     temp_file.write(file_stream.read())  # Write the BytesIO stream to the temp file
        #     temp_file_path = temp_file.name  # Get the temporary file path

        # # Extract text from the .docx file using python-docx
        # doc = Document(temp_file_path)
        # text = '\n'.join([para.text for para in doc.paragraphs])

        # return text
        document = docx.Document(file_stream)
        # Extract text from the .docx file using python-docx
        text = '\n'.join([para.text for para in document.paragraphs])
        return text

    def create_file(self, text: str) -> BytesIO:
        document = docx.Document()
        document.add_paragraph(text)
        docx_stream = BytesIO()
        document.save(docx_stream)
        docx_stream.seek(0)

        return docx_stream
