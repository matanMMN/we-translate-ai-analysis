"""
Setup instructions:
pip install openai PyMuPDF pillow python-docx
"""
from fastapi import UploadFile
from typing import Tuple
from openai import OpenAI
import fitz 
import io
import base64
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image  

def convert_pdf_bytes_to_images(pdf_bytes: bytes):
    """Convert PDF pages to a list of images using PyMuPDF"""
    pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
    images = []
    
    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        zoom = 2
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        images.append(img)
    
    pdf_document.close()
    return images


def batch_process_images(client, images, batch_size=3):
    """Process multiple images in a single API call"""
    all_content = []
    
    for i in range(0, len(images), batch_size):
        batch = images[i:i + batch_size]
        batch_messages = []
        
        for img in batch:
            base64_img = encode_image_object(img)
            batch_messages.append({
                "type": "image_url",
                "image_url": {"url": f"data:image/png;base64,{base64_img}"}
            })
        
        response = client.chat.completions.create(
            model='gpt-4o',
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": """For each image, extract the text exactly as it appears, maintaining:
                        1. Text formatting (bold, italics, etc.) - mark using markdown
                        2. Section structure and hierarchy
                        3. Lists and bullet points
                        4. Any special characters or symbols
                        5. Tables (if any)

                        Separate each page with '===PAGE_BREAK==='."""
                        },
                        *batch_messages
                    ]
                }
            ],
            max_tokens=4096,
        )
        
        content = response.choices[0].message.content
        pages = [page.strip() for page in content.split('===PAGE_BREAK===')]
        # Filter out empty pages
        pages = [page for page in pages if page.strip()]
        all_content.extend(pages)
    
    return all_content


def encode_image_object(image):
    """Encode a PIL Image object to base64"""
    img_byte_arr = io.BytesIO()
    image.save(img_byte_arr, format='PNG')
    img_byte_arr = img_byte_arr.getvalue()
    return base64.b64encode(img_byte_arr).decode('utf-8')

def create_docx_from_markdown(markdown_content, doc):
    """Convert markdown content to DOCX formatting"""

    markdown_content = markdown_content.replace('```', '') 
    markdown_content = markdown_content.replace('markdown', '')
    markers_to_remove = [
        "markdown"
        "===",
        "Text:",
        "Image Description:",
        "Extracted Text:",
        "PAGE_BREAK"
    ]
    lines = markdown_content.split('\n')
    
    for line in lines:
        # Handle headers

        if (not line.strip() or 
            line.strip() == '`' or 
            any(marker in line for marker in markers_to_remove)):
            continue
        
        if line.startswith('#'):
            level = line.count('#')
            text = line.strip('#').strip()
            doc.add_heading(text, level=level)
            continue
            
        # Handle bullet points
        if line.strip().startswith('*') or line.strip().startswith('-'):
            text = line.strip('*- ').strip()
            doc.add_paragraph(text, style='List Bullet')
            continue
            
        # Handle numbered lists
        if line.strip() and line.strip()[0].isdigit() and '. ' in line:  # Fixed this line
            text = line.split('.', 1)[1].strip()
            doc.add_paragraph(text, style='List Number')
            continue
            
        # Handle bold and italic text
        if '**' in line or '*' in line or '_' in line:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1:  # Bold text
                    p.add_run(part).bold = True
                else:
                    # Handle italic text in non-bold parts
                    italic_parts = part.split('*')
                    for j, ipart in enumerate(italic_parts):
                        if j % 2 == 1:  # Italic text
                            p.add_run(ipart).italic = True
                        else:
                            p.add_run(ipart)
            continue
            
        # Regular text
        if line.strip():
            doc.add_paragraph(line.strip())

async def pdf_to_docx_bytes(pdf_file: UploadFile, api_key: str) -> Tuple[bytes, str]:
    """Convert PDF to DOCX while maintaining structure and formatting"""
    client = OpenAI(api_key=api_key)    
    pdf_bytes = await pdf_file.read()
    images = convert_pdf_bytes_to_images(pdf_bytes)
    doc = Document()
    markdown_contents = batch_process_images(client, images)

    for i, markdown_content in enumerate(markdown_contents):
        if i > 0:  # Only add page break after the first page
            doc.add_page_break()
        print(f"Converting page {i+1} to DOCX format...")
        create_docx_from_markdown(markdown_content.strip(), doc)
    
    # Save to bytes buffer instead of file
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    # Generate new filename
    original_filename = pdf_file.filename
    new_filename = original_filename.rsplit('.', 1)[0] + '.docx'
    
    return docx_buffer.getvalue(), new_filename